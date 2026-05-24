from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "entregables" / "Informe_Final_SIDERAE_Blenkir_Preliminar.docx"
MOCKUPS = ROOT / "docs" / "ui" / "mockups"


TITLE = "Sistema Inteligente de Deteccion Temprana de Riesgo Academico y Desercion Estudiantil"
PROJECT = "SIDERAE-Blenkir"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    if table.rows:
        set_repeat_table_header(table.rows[0])
        for cell in table.rows[0].cells:
            set_cell_shading(cell, "F05A0E")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(3)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color="FFFFFF")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    style_table(table)
    doc.add_paragraph()
    return table


def add_placeholder(doc, label):
    p = doc.add_paragraph()
    p.style = "Intense Quote"
    r = p.add_run(f"PENDIENTE / ESPACIO EN BLANCO: {label}")
    r.bold = True
    r.font.color.rgb = RGBColor(180, 0, 0)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.italic = True
        r.font.size = Pt(9)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(11)

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Arial"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        styles[name].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(12)


def cover(doc):
    for text in [
        "UNIVERSIDAD CONTINENTAL",
        "FACULTAD DE INGENIERIA",
        "ESCUELA ACADEMICO PROFESIONAL DE INGENIERIA DE SISTEMAS E INFORMATICA",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(12)

    doc.add_paragraph("\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PROYECTO")
    r.bold = True
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'"{TITLE} - {PROJECT}"')
    r.bold = True
    r.font.size = Pt(15)

    doc.add_paragraph("\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("PRESENTADO POR:").bold = True

    add_table(
        doc,
        ["APELLIDOS Y NOMBRES", "CODIGO"],
        [
            ["Diego Carhuamaca Vasquez", ""],
            ["Ernesto Chuchon Sotelo", ""],
            ["", ""],
        ],
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("ASESOR:").bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("DR. MAGLIONI ARANA CAPARACHIN")

    doc.add_paragraph("\n\n")
    for text in ["HUANCAYO - PERU", "2026"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text).bold = True

    add_placeholder(doc, "Completar codigos universitarios, NRC/seccion y validar nombres oficiales.")
    doc.add_page_break()


def toc(doc):
    add_heading(doc, "CONTENIDO GENERAL", 1)
    doc.add_paragraph("Tabla de contenido preliminar. Actualizar en Microsoft Word con Referencias > Actualizar tabla cuando el documento final este cerrado.")
    sections = [
        "CAPITULO 1. INFORMACION GENERAL DEL PROYECTO",
        "CAPITULO 2. CONTEXTO ORGANIZACIONAL Y ANALISIS DEL PROBLEMA",
        "CAPITULO 3. ANALISIS DE PROCESOS DE NEGOCIO",
        "CAPITULO 4. ANALISIS DE REQUERIMIENTOS DEL SISTEMA",
        "CAPITULO 5. PLANIFICACION DEL PROYECTO Y PLAN DE CALIDAD",
        "CAPITULO 6. DISENO DEL SISTEMA",
        "CAPITULO 7. ARQUITECTURA TECNOLOGICA DEL SISTEMA",
        "CAPITULO 8. DESARROLLO DEL SISTEMA",
        "CAPITULO 9. CONTROL DE VERSIONES Y GESTION DEL REPOSITORIO",
        "CAPITULO 10. DOCKERIZACION Y DESPLIEGUE DE MODULOS",
        "CAPITULO 11. ESTRATEGIA DE PRUEBAS DE SOFTWARE",
        "CAPITULO 12. AUTOMATIZACION DE PRUEBAS",
        "CAPITULO 13. METRICAS DE CALIDAD",
        "CAPITULO 14. IMPLEMENTACION Y MONITOREO",
        "CONCLUSIONES",
        "RECOMENDACIONES",
        "REFERENCIAS",
        "ANEXOS",
    ]
    add_bullets(doc, sections)
    doc.add_page_break()


def chapter_1(doc):
    add_heading(doc, "CAPITULO 1", 1)
    add_heading(doc, "INFORMACION GENERAL DEL PROYECTO", 1)
    add_heading(doc, "1.1. Resumen Ejecutivo", 2)
    doc.add_paragraph(
        "SIDERAE-Blenkir es un sistema web orientado a la deteccion temprana de riesgo academico y desercion estudiantil en el Colegio Blenkir. "
        "El prototipo integra gestion de estudiantes, registro de notas, asistencia y variables socioeconomicas, calculo de riesgo academico, generacion de alertas, intervenciones, dashboard y exportacion basica de reportes."
    )
    doc.add_paragraph(
        "La solucion se implementa con una arquitectura desacoplada: frontend React, backend Laravel API REST, base de datos MySQL, microservicio Flask para calculo de riesgo y orquestacion local mediante Docker Compose. "
        "El alcance actual corresponde a un prototipo academico funcional con diferencias explicitas frente al DRS formal; por ello, algunos requerimientos se documentan como parciales o pendientes."
    )
    add_heading(doc, "1.2. Introduccion", 2)
    doc.add_paragraph(
        "El proyecto se ubica en el campo de la ingenieria de software aplicada a contextos educativos. Su proposito es apoyar la identificacion temprana de estudiantes con posibles senales de riesgo academico, permitiendo que docentes, coordinadores, psicologos o directivos actuen de manera preventiva."
    )
    doc.add_paragraph(
        "El informe documenta el contexto organizacional, el analisis del proceso, los requerimientos, la planificacion, el diseno, la arquitectura, el desarrollo por iteraciones, la gestion del repositorio, la dockerizacion, las pruebas y el cierre de calidad del prototipo."
    )


def chapter_2(doc):
    add_heading(doc, "CAPITULO 2", 1)
    add_heading(doc, "CONTEXTO ORGANIZACIONAL Y ANALISIS DEL PROBLEMA", 1)
    add_heading(doc, "2.1. Contexto de la Organizacion", 2)
    doc.add_paragraph(
        "El escenario de aplicacion corresponde al Colegio Blenkir, institucion educativa privada ubicada en Huancayo, Junin, con sedes referenciadas en la documentacion del proyecto como Chilca y Auquimarca. "
        "El sistema se enfoca en apoyar procesos de seguimiento academico, asistencia, registro socioeconomico, analisis de riesgo y gestion de alertas preventivas."
    )
    add_table(
        doc,
        ["Elemento", "Descripcion preliminar"],
        [
            ["Organizacion", "Colegio Blenkir"],
            ["Area del problema", "Seguimiento academico, asistencia, riesgo y desercion estudiantil"],
            ["Usuarios principales", "Administrador, docente, coordinador academico, psicologo/tutor y directivo"],
            ["Herramientas actuales", "Pendiente de documentar formalmente; se asume gestion parcialmente manual y hojas de calculo/documentos internos."],
        ],
    )
    add_placeholder(doc, "Agregar descripcion institucional oficial, organigrama si corresponde y herramientas reales usadas antes del sistema.")

    add_heading(doc, "2.2. Identificacion del Problema", 2)
    doc.add_paragraph(
        "El problema central identificado es la dificultad para detectar oportunamente estudiantes con riesgo academico o posible desercion cuando la informacion se encuentra dispersa, incompleta o se revisa de manera reactiva. "
        "La ausencia de integracion entre notas, asistencia, variables socioeconomicas, reportes conductuales y acciones de intervencion limita la capacidad preventiva de la institucion."
    )
    add_bullets(
        doc,
        [
            "Informacion academica y socioeconomica dispersa entre registros manuales o digitales no integrados.",
            "Dificultad para priorizar estudiantes segun nivel de riesgo.",
            "Seguimiento tardio de alertas e intervenciones.",
            "Trazabilidad parcial de acciones realizadas por usuarios.",
            "Necesidad de dashboard y reportes para toma de decisiones.",
        ],
    )
    add_placeholder(doc, "Completar causas y consecuencias con evidencia institucional o entrevistas.")


def chapter_3(doc):
    add_heading(doc, "CAPITULO 3", 1)
    add_heading(doc, "ANALISIS DE PROCESOS DE NEGOCIO", 1)
    add_heading(doc, "3.1. Descripcion del Proceso Actual", 2)
    doc.add_paragraph(
        "De forma preliminar, el proceso actual puede describirse como un flujo reactivo: el docente registra o revisa informacion academica, identifica casos preocupantes, comunica la situacion a coordinacion o direccion y se generan acciones de apoyo segun disponibilidad de datos. "
        "La informacion no siempre se encuentra consolidada en un unico sistema."
    )
    add_numbered(
        doc,
        [
            "Registro de notas, asistencia u observaciones por parte de docentes.",
            "Revision manual de casos con bajo rendimiento o inasistencias.",
            "Comunicacion informal o documentada a coordinacion/directivo.",
            "Derivacion o intervencion cuando el caso ya es visible.",
            "Seguimiento posterior con trazabilidad variable.",
        ],
    )
    add_heading(doc, "3.2. Modelado del Proceso Actual (AS-IS)", 2)
    add_placeholder(doc, "Insertar diagrama BPMN AS-IS.")
    add_heading(doc, "3.3. Problemas del Proceso Actual", 2)
    add_bullets(
        doc,
        [
            "Procesos manuales con alto esfuerzo de revision.",
            "Duplicacion de registros y falta de integracion.",
            "Riesgo de intervenciones tardias.",
            "Dificultad para obtener indicadores consolidados por aula, sede o nivel.",
            "Falta de trazabilidad completa entre alerta, intervencion y cierre.",
        ],
    )
    add_heading(doc, "3.4. Modelado del Proceso Propuesto (TO-BE)", 2)
    doc.add_paragraph(
        "El proceso propuesto incorpora SIDERAE-Blenkir como plataforma central para registrar datos, procesar riesgo, generar alertas y documentar intervenciones. "
        "El flujo busca pasar de una gestion reactiva a una gestion preventiva basada en informacion integrada."
    )
    add_numbered(
        doc,
        [
            "El usuario autenticado registra datos academicos, asistencia y variables socioeconomicas.",
            "El backend valida permisos y almacena informacion en MySQL.",
            "El sistema procesa riesgo mediante integracion Laravel-Flask.",
            "Se genera o actualiza el indice de riesgo y, si corresponde, la alerta.",
            "El usuario autorizado registra intervenciones y cierre.",
            "Dashboard y reportes apoyan el seguimiento institucional.",
        ],
    )
    add_placeholder(doc, "Insertar diagrama BPMN TO-BE.")


def chapter_4(doc):
    add_heading(doc, "CAPITULO 4", 1)
    add_heading(doc, "ANALISIS DE REQUERIMIENTOS DEL SISTEMA", 1)
    add_heading(doc, "4.1. Identificacion de Actores del Sistema", 2)
    add_table(
        doc,
        ["Actor", "Descripcion"],
        [
            ["Administrador", "Gestiona usuarios, roles, permisos, materias y configuraciones principales."],
            ["Docente", "Registra y consulta datos academicos, asistencia, estudiantes, alertas e intervenciones segun permisos."],
            ["Coordinador academico", "Gestiona informacion academica y procesa riesgo segun permisos definidos."],
            ["Psicologo/Tutor", "Consulta alertas y registra intervenciones o seguimiento de casos."],
            ["Directivo", "Consulta dashboard, alertas y reportes para toma de decisiones institucionales."],
            ["Sistema", "Calcula riesgo, genera alertas y registra trazabilidad automatica."],
        ],
    )
    add_heading(doc, "4.2. Requerimientos Funcionales", 2)
    rf_rows = [
        ["RF-01", "Carga e importacion de datos academicos", "Parcial"],
        ["RF-02", "Registro digital de asistencia semanal", "Confirmado"],
        ["RF-03", "Importacion de resultados del Fast Test", "Pendiente/verificar"],
        ["RF-04", "Registro digital de reportes conductuales", "Pendiente/verificar"],
        ["RF-05", "Integracion de variables socioeconomicas", "Confirmado"],
        ["RF-06", "Procesamiento multivariable y calculo de indice de riesgo", "Parcial/confirmado en flujo"],
        ["RF-07", "Evaluacion automatica del nivel de riesgo", "Confirmado"],
        ["RF-08", "Emision de alertas tempranas accionables", "Confirmado"],
        ["RF-09", "Intervencion preventiva del docente", "Confirmado"],
        ["RF-10", "Decision de derivacion por directivo", "Pendiente/verificar"],
        ["RF-11", "Atencion psicologica preventiva con perfil integrado", "Pendiente/verificar"],
        ["RF-12", "Comunicacion formal y trazable con familia", "Pendiente"],
        ["RF-13", "Registro de accion tomada y cierre de alerta", "Confirmado/parcial"],
        ["RF-14", "Panel de visualizacion de riesgo", "Parcial"],
        ["RF-15", "Gestion de usuarios y control de acceso por rol", "Confirmado"],
        ["RF-16", "Exportacion de reportes PDF", "Parcial"],
        ["RF-17", "Registro de auditoria de acciones", "Parcial"],
        ["RF-18", "Reentrenamiento del modelo ML", "Pendiente"],
        ["RF-19", "Semaforo de completitud de datos", "Pendiente"],
        ["RF-20", "Historial de riesgo por estudiante", "Parcial"],
    ]
    add_table(doc, ["ID", "Requerimiento", "Estado preliminar"], rf_rows)
    add_heading(doc, "4.3. Requerimientos no Funcionales", 2)
    add_table(
        doc,
        ["ID", "Categoria", "Descripcion resumida"],
        [
            ["RNF-01", "Rendimiento", "Dashboard menor a 3 s con datos procesados; ML hasta 10 s en background."],
            ["RNF-02", "Disponibilidad", "Disponibilidad objetivo en horario escolar; recuperacion maxima definida en DRS."],
            ["RNF-03", "Seguridad", "HTTPS/TLS, bcrypt, control de acceso por Spatie Permission."],
            ["RNF-04", "Usabilidad", "Interfaz responsiva desde 768 px y flujos principales acotados."],
            ["RNF-05", "Mantenibilidad", "Buenas practicas Laravel/React y pruebas en modulos criticos."],
            ["RNF-06", "Escalabilidad", "Arquitectura desacoplada y ML reemplazable."],
            ["RNF-07", "Trazabilidad", "Registro de acciones relevantes en activity_log."],
            ["RNF-08", "Compatibilidad", "Chrome y Firefox modernos."],
            ["RNF-09", "Portabilidad", "Despliegue reproducible con Docker Compose."],
            ["RNF-10", "Integridad", "Validacion de datos criticos antes de procesar ML."],
        ],
    )
    add_heading(doc, "4.4. Casos de Uso del Sistema", 2)
    add_placeholder(doc, "Insertar diagrama de casos de uso.")
    add_bullets(
        doc,
        [
            "Iniciar sesion y consultar permisos.",
            "Gestionar estudiantes.",
            "Registrar notas, asistencia y variables socioeconomicas.",
            "Procesar riesgo academico.",
            "Consultar y atender alertas.",
            "Registrar intervenciones y cerrar alertas.",
            "Consultar dashboard y exportar reporte PDF.",
        ],
    )


def chapter_5(doc):
    add_heading(doc, "CAPITULO 5", 1)
    add_heading(doc, "PLANIFICACION DEL PROYECTO Y PLAN DE CALIDAD", 1)
    add_heading(doc, "5.1. Alcance del Proyecto", 2)
    add_bullets(
        doc,
        [
            "Incluye autenticacion, roles y permisos.",
            "Incluye gestion de estudiantes, materias, notas, asistencia y variables socioeconomicas.",
            "Incluye calculo de riesgo academico mediante Laravel y Flask.",
            "Incluye alertas, intervenciones, cierre de alertas, dashboard y export PDF basico.",
            "Excluye por ahora reentrenamiento ML completo, semaforo de completitud, comunicacion familiar formal y algunos reportes avanzados.",
        ],
    )
    add_heading(doc, "5.2. Herramientas Tecnologicas del Proyecto", 2)
    add_table(
        doc,
        ["Componente", "Tecnologia"],
        [
            ["Frontend", "React 18, Vite, Tailwind CSS"],
            ["Backend", "Laravel, PHP 8.3, Laravel Sanctum"],
            ["Base de datos", "MySQL 8"],
            ["ML Service", "Python Flask"],
            ["Autorizacion", "Spatie Laravel Permission"],
            ["Auditoria", "Spatie Activitylog"],
            ["Reportes", "Barryvdh DomPDF"],
            ["Infraestructura", "Docker y Docker Compose"],
            ["Control de versiones", "Git y GitHub"],
        ],
    )
    add_heading(doc, "5.3. Normas y Estandares de Calidad", 2)
    doc.add_paragraph(
        "El proyecto utiliza normas y estandares como referencia academica y orientativa, sin afirmar certificacion formal. Se consideran ISO/IEC 25010 para calidad del producto, ISO/IEC 29119 para pruebas de software, ISO 9001 para trazabilidad de procesos y buenas practicas de seguridad de la informacion."
    )
    add_heading(doc, "5.4. Plan de Pruebas del Proyecto", 2)
    doc.add_paragraph(
        "El plan de pruebas documenta niveles unitarios, integracion, sistema, funcionales y aceptacion. Para backend se emplean pruebas PHPUnit/Feature Tests; para frontend se plantea build y pruebas E2E cuando exista suite estable. Las pruebas se organizan por requerimiento funcional y por sprint."
    )
    add_heading(doc, "5.5. Lineamientos de Seguridad Informatica", 2)
    add_bullets(
        doc,
        [
            "Autenticacion por Laravel Sanctum.",
            "Autorizacion backend mediante Spatie Permission.",
            "No exponer archivos .env ni credenciales reales.",
            "Validacion de entradas mediante Form Requests.",
            "Registro de acciones criticas mediante activity_log.",
            "Separacion de responsabilidades entre frontend, backend, base de datos y ML service.",
        ],
    )


def chapter_6(doc):
    add_heading(doc, "CAPITULO 6", 1)
    add_heading(doc, "DISENO DEL SISTEMA", 1)
    add_heading(doc, "6.1. Arquitectura Conceptual del Sistema", 2)
    doc.add_paragraph(
        "La arquitectura conceptual se organiza en cuatro capas: interfaz de usuario, API de negocio, persistencia y servicio de calculo de riesgo. El usuario interactua con React, React consume Laravel por HTTP, Laravel persiste en MySQL y coordina el procesamiento de riesgo con Flask."
    )
    add_table(
        doc,
        ["Capa", "Responsabilidad"],
        [
            ["Frontend React", "Interfaz de usuario, consumo API, estados por permisos."],
            ["Backend Laravel", "Reglas de negocio, autenticacion, autorizacion, APIs y auditoria."],
            ["MySQL", "Persistencia de usuarios, estudiantes, datos academicos, alertas e indices."],
            ["Flask ML", "Calculo de indice y nivel de riesgo academico."],
        ],
    )
    add_placeholder(doc, "Insertar diagrama conceptual de arquitectura.")
    add_heading(doc, "6.2. Modelo UML del Sistema", 2)
    add_placeholder(doc, "Insertar diagrama de clases UML.")
    add_placeholder(doc, "Insertar diagrama de secuencia del procesamiento de riesgo.")
    add_heading(doc, "6.3. Diseno de Interfaces de Usuario", 2)
    doc.add_paragraph(
        "La interfaz se disena sobre una guia visual institucional con paleta naranja, tarjetas blancas, fondo claro, navegacion lateral y componentes reutilizables. La carpeta de mockups contiene 12 pantallas de referencia."
    )
    mockup_files = [
        "01-login.png",
        "02-dashboard.png",
        "03-listado-estudiantes.png",
        "04-registro-edicion-estudiante.png",
        "05-perfil-estudiante.png",
        "06-registro-notas.png",
        "07-registro-asistencia.png",
        "08-variables-socioeconomicas.png",
        "09-riesgo-academico.png",
        "10-listado-alertas.png",
        "11-detalle-alerta.png",
        "12-registro-intervencion.png",
    ]
    for i, name in enumerate(mockup_files, start=1):
        path = MOCKUPS / name
        if path.exists():
            try:
                doc.add_picture(str(path), width=Inches(5.8))
                add_caption(doc, f"Figura {i}. Mockup de interfaz: {name}")
            except Exception:
                add_placeholder(doc, f"No se pudo insertar mockup {name}.")
    add_heading(doc, "6.4. Diseno de Base de Datos", 2)
    doc.add_paragraph(
        "La base de datos se construye mediante migraciones Laravel y modelos Eloquent. Las entidades principales detectadas son User, Estudiante, Materia, Nota, Asistencia, VariableSocioeconomica, IndiceRiesgo, Alerta, Intervencion y ReporteConductual."
    )
    add_table(
        doc,
        ["Entidad", "Uso principal"],
        [
            ["users", "Usuarios del sistema y autenticacion."],
            ["estudiantes", "Datos personales y academicos base."],
            ["materias", "Catalogo institucional por sede/nivel/grado/anio."],
            ["notas", "Calificaciones por estudiante y materia."],
            ["asistencias", "Registros de asistencia."],
            ["variables_socioeconomicas", "Variables de contexto familiar/social."],
            ["indices_riesgo", "Historial de indices y niveles de riesgo."],
            ["alertas", "Alertas tempranas generadas o gestionadas."],
            ["intervenciones", "Acciones preventivas realizadas."],
            ["activity_log", "Auditoria de acciones criticas."],
        ],
    )
    add_placeholder(doc, "Insertar modelo entidad-relacion / DER.")


def chapter_7(doc):
    add_heading(doc, "CAPITULO 7", 1)
    add_heading(doc, "ARQUITECTURA TECNOLOGICA DEL SISTEMA", 1)
    add_heading(doc, "7.1. Tecnologias del Frontend", 2)
    doc.add_paragraph("El frontend utiliza React con Vite y Tailwind CSS. Sus componentes principales incluyen LoginForm, DashboardPanel, EstudiantesPanel, MateriasPanel, AlertasPanel y componentes UI reutilizables.")
    add_heading(doc, "7.2. Tecnologias del Backend", 2)
    doc.add_paragraph("El backend Laravel expone API REST, gestiona autenticacion con Sanctum, autorizacion con Spatie Permission, auditoria con Spatie Activitylog y reportes PDF con DomPDF.")
    add_heading(doc, "7.3. Base de Datos del Sistema", 2)
    doc.add_paragraph("MySQL 8 es la base de datos principal. Laravel administra la estructura mediante migraciones, seeders y modelos Eloquent.")
    add_heading(doc, "7.4. Infraestructura de Desarrollo", 2)
    add_table(
        doc,
        ["Servicio", "Puerto host", "Descripcion"],
        [
            ["app-frontend", "5173", "Aplicacion React/Vite"],
            ["app-backend", "8000", "API Laravel"],
            ["ml-engine", "5000", "Servicio Flask"],
            ["db-mysql", "3307", "Base de datos MySQL"],
        ],
    )


def chapter_8(doc):
    add_heading(doc, "CAPITULO 8", 1)
    add_heading(doc, "DESARROLLO DEL SISTEMA", 1)
    add_heading(doc, "8.1. Iteracion 1: Configuracion Inicial del Proyecto", 2)
    add_bullets(doc, ["Configuracion Docker y servicios base.", "Estructura backend, frontend, ML service y base de datos.", "Health checks y verificacion de arranque."])
    add_heading(doc, "8.2. Iteracion 2: Desarrollo de Funcionalidades Basicas", 2)
    add_bullets(doc, ["Login React-Laravel.", "Endpoint /api/me.", "Roles y permisos minimos.", "CRUD de estudiantes."])
    add_heading(doc, "8.3. Iteracion 3: Implementacion de Modulos Funcionales", 2)
    add_bullets(doc, ["Notas, asistencia y variables socioeconomicas.", "Procesamiento de riesgo.", "Alertas e intervenciones.", "Dashboard y export PDF basico."])
    add_heading(doc, "8.4. Iteraciones Posteriores", 2)
    add_table(
        doc,
        ["Sprint", "Resultado documentado"],
        [
            ["Sprint 4", "Integracion Laravel-Flask para riesgo academico."],
            ["Sprint 5", "Alertas, intervenciones y cierre."],
            ["Sprint 6A/6B", "Dashboard, filtros y export PDF parcial."],
            ["Sprint 7A/7B", "Redisenio UI y navegacion segun mockups."],
            ["Sprint 7.6A/7.6B", "Materias, notas masivas y asistencia masiva."],
            ["Sprint 8", "Seguridad, roles, auditoria y control de accesos."],
            ["Sprint 9", "Pruebas integrales y regresion."],
            ["Sprint 10", "Documentacion final y cierre de calidad."],
        ],
    )


def chapter_9(doc):
    add_heading(doc, "CAPITULO 9", 1)
    add_heading(doc, "CONTROL DE VERSIONES Y GESTION DEL REPOSITORIO", 1)
    add_heading(doc, "9.1. Repositorio del Proyecto", 2)
    doc.add_paragraph("Repositorio remoto: https://github.com/Keterod/siderae-blenkir.git")
    add_heading(doc, "9.2. Estrategia de Control de Versiones", 2)
    doc.add_paragraph("Se utiliza Git como sistema de control de versiones. Los commits reflejan incrementos por sprint, documentacion, funcionalidades, pruebas y ajustes de arquitectura.")
    add_heading(doc, "9.3. Gestion de Ramas del Proyecto", 2)
    add_bullets(doc, ["Rama principal: main.", "Ramas auxiliares detectadas: sprint-6, gg.", "La estrategia formal de ramas por feature esta pendiente de documentar si se requiere mayor rigor."])
    add_heading(doc, "9.4. Registro de Commits Relevantes", 2)
    add_table(
        doc,
        ["Commit", "Descripcion"],
        [
            ["a63e4e6", "Actualizacion de estado del arte."],
            ["65090e0", "Flujo academico masivo y control de accesos sprint 8."],
            ["cffdd38", "Sprint 7.5A auditoria y documentacion tecnica."],
            ["2c51c0d", "Dashboard funcional Sprint 6A."],
            ["abdfaf5", "Sprint 5 alertas e intervenciones."],
            ["8590b2f", "Sprint 4 integracion ML Laravel + Flask."],
            ["7b02e62", "Sprint 3A CRUD estudiantes."],
            ["8d0504c", "Sprint 2 autenticacion, roles y permisos."],
            ["ccfa576", "Sprint 1 Docker setup y health checks."],
        ],
    )


def chapter_10(doc):
    add_heading(doc, "CAPITULO 10", 1)
    add_heading(doc, "DOCKERIZACION Y DESPLIEGUE DE MODULOS", 1)
    add_heading(doc, "10.1. Introduccion a Docker en el Proyecto", 2)
    doc.add_paragraph("Docker permite ejecutar el prototipo con servicios separados y configuracion reproducible para frontend, backend, base de datos y ML service.")
    add_heading(doc, "10.2. Dockerizacion del Backend", 2)
    doc.add_paragraph("El backend Laravel se ejecuta en el servicio app-backend, expuesto por el puerto 8000 y conectado a MySQL y al servicio ML mediante variables de entorno.")
    add_heading(doc, "10.3. Dockerizacion del Frontend", 2)
    doc.add_paragraph("El frontend React/Vite se ejecuta en app-frontend, expuesto por el puerto 5173 y configurado para consumir la API Laravel mediante VITE_API_URL.")
    add_heading(doc, "10.4. Orquestacion con Docker Compose", 2)
    add_table(
        doc,
        ["Servicio", "Contenedor", "Responsabilidad"],
        [
            ["db-mysql", "siderae_mysql", "Base de datos MySQL"],
            ["app-backend", "siderae_backend", "API Laravel"],
            ["app-frontend", "siderae_frontend", "Interfaz React"],
            ["ml-engine", "siderae_ml", "Calculo de riesgo Flask"],
        ],
    )


def chapter_11(doc):
    add_heading(doc, "CAPITULO 11", 1)
    add_heading(doc, "ESTRATEGIA DE PRUEBAS DE SOFTWARE", 1)
    add_heading(doc, "11.1. Enfoque de Pruebas del Proyecto", 2)
    doc.add_paragraph("El enfoque combina pruebas automatizadas de backend, pruebas funcionales manuales y validacion incremental por sprints. El plan de pruebas usa como referencia el Modelo en V y un enfoque TDD en los modulos donde aplica.")
    add_heading(doc, "11.2. Niveles de Pruebas Aplicados", 2)
    add_bullets(doc, ["Pruebas unitarias.", "Pruebas de integracion.", "Pruebas funcionales.", "Pruebas de sistema.", "Pruebas de aceptacion preliminar."])
    add_heading(doc, "11.3. Tipos de Pruebas Ejecutadas", 2)
    add_bullets(doc, ["Funcionalidad.", "Validacion de datos.", "Interfaz de usuario.", "Seguridad y permisos.", "Auditoria y trazabilidad.", "Regresion."])
    add_heading(doc, "11.4. Plan de Ejecucion de Pruebas", 2)
    add_table(
        doc,
        ["Modulo", "Evidencia/Comando esperado"],
        [
            ["Backend Laravel", "php artisan test"],
            ["Frontend React", "npm run build"],
            ["Flujos E2E", "Cypress si se implementa suite estable"],
            ["Docker", "docker compose ps / logs"],
            ["Base de datos", "Validacion de seeders, migraciones y conteos demo"],
        ],
    )
    add_placeholder(doc, "Insertar resultados reales finales de ejecucion de pruebas.")


def chapter_12(doc):
    add_heading(doc, "CAPITULO 12", 1)
    add_heading(doc, "AUTOMATIZACION DE PRUEBAS", 1)
    add_heading(doc, "12.1. Herramientas de Automatizacion", 2)
    add_bullets(doc, ["PHPUnit / Laravel Feature Tests.", "Factories y seeders Laravel.", "Http::fake para aislar llamadas al microservicio ML.", "Cypress recomendado para flujos criticos E2E.", "npm run build para verificacion de frontend."])
    add_heading(doc, "12.2. Configuracion del Entorno de Pruebas", 2)
    doc.add_paragraph("Las pruebas se ejecutan preferentemente en entorno Docker, con contenedores para Laravel, React, MySQL y Flask. El README documenta comandos de arranque, seeders demo y validacion de conteos.")
    add_heading(doc, "12.3. Scripts de Pruebas Automatizadas", 2)
    add_table(
        doc,
        ["Archivo/Modulo", "Proposito"],
        [
            ["backend/tests/Feature/EstudianteTest.php", "Pruebas CRUD estudiantes."],
            ["backend/tests/Feature/DatosAcademicosTest.php", "Pruebas de notas, asistencia y variables."],
            ["backend/tests/Feature/RiesgoTest.php", "Pruebas de procesamiento de riesgo."],
            ["backend/tests/Feature/AlertaIntervencionTest.php", "Pruebas de alertas e intervenciones."],
            ["backend/tests/Feature/DashboardTest.php", "Pruebas de dashboard/export."],
            ["backend/tests/Feature/ActivityLogTest.php", "Pruebas de auditoria."],
        ],
    )
    add_heading(doc, "12.4. Ejecucion Automatica de Pruebas", 2)
    add_placeholder(doc, "Insertar salida final de php artisan test, npm run build y Cypress si aplica.")


def chapter_13(doc):
    add_heading(doc, "CAPITULO 13", 1)
    add_heading(doc, "METRICAS DE CALIDAD", 1)
    add_heading(doc, "13.1. Ejecucion de Casos de Pruebas", 2)
    add_table(
        doc,
        ["Indicador", "Valor preliminar"],
        [
            ["Casos definidos", "Pendiente de consolidar desde fichas manuales y automatizadas"],
            ["Casos ejecutados", ""],
            ["Casos aprobados", ""],
            ["Casos fallidos", ""],
            ["Porcentaje de exito", ""],
        ],
    )
    add_heading(doc, "13.2. Registro de Defectos", 2)
    add_table(doc, ["ID", "Descripcion", "Severidad", "Estado", "Accion correctiva"], [["", "", "", "", ""], ["", "", "", "", ""]])
    add_heading(doc, "13.3. Metricas de Calidad del Software", 2)
    add_bullets(doc, ["Tasa de defectos.", "Porcentaje de pruebas aprobadas.", "Cobertura de pruebas en modulos criticos.", "Estabilidad del entorno Docker.", "Trazabilidad de acciones en activity_log."])
    add_heading(doc, "13.4. Evaluacion de Calidad basada en Estandares", 2)
    doc.add_paragraph("La evaluacion se plantea como alineacion orientativa a ISO/IEC 25010, ISO/IEC 29119 e ISO 9001, sin declarar certificacion ni auditoria externa.")
    add_placeholder(doc, "Completar matriz de evaluacion ISO 25010 con valores finales.")


def chapter_14(doc):
    add_heading(doc, "CAPITULO 14", 1)
    add_heading(doc, "IMPLEMENTACION Y MONITOREO", 1)
    add_heading(doc, "14.1. Preparacion del Entorno de Implementacion", 2)
    doc.add_paragraph("El entorno se prepara copiando archivos .env.example, levantando Docker Compose, generando APP_KEY si corresponde, ejecutando migraciones/seeders y verificando servicios.")
    add_heading(doc, "14.2. Implementacion del Sistema", 2)
    add_numbered(doc, ["Clonar repositorio.", "Crear archivos .env.", "Ejecutar docker compose up -d --build.", "Generar APP_KEY si falta.", "Ejecutar migraciones y seeders.", "Abrir frontend en localhost:5173."])
    add_heading(doc, "14.3. Verificacion de Funcionamiento", 2)
    add_bullets(doc, ["Frontend: http://localhost:5173.", "Backend: http://localhost:8000.", "ML Service: http://localhost:5000.", "MySQL host: localhost:3307."])
    add_placeholder(doc, "Insertar capturas reales del sistema levantado y docker compose ps.")
    add_heading(doc, "14.4. Monitoreo del Sistema", 2)
    doc.add_paragraph("El monitoreo preliminar se realiza mediante logs de Docker, respuestas HTTP de servicios, pruebas automatizadas y registros activity_log en acciones criticas. No se documenta aun una plataforma externa de monitoreo.")
    add_placeholder(doc, "Completar actividades de monitoreo final y evidencias.")


def closing(doc):
    add_heading(doc, "CONCLUSIONES", 1)
    add_numbered(
        doc,
        [
            "El prototipo SIDERAE-Blenkir integra modulos clave para seguimiento academico, procesamiento de riesgo, alertas e intervenciones.",
            "La arquitectura desacoplada con React, Laravel, MySQL, Flask y Docker facilita la ejecucion local y la separacion de responsabilidades.",
            "El proyecto presenta avances significativos, aunque mantiene requerimientos parciales o pendientes frente al DRS formal.",
            "La calidad del sistema se apoya en pruebas backend, control de permisos, auditoria parcial y documentacion tecnica por sprints.",
        ],
    )
    add_placeholder(doc, "Ajustar conclusiones con resultados finales de pruebas y validacion.")
    add_heading(doc, "RECOMENDACIONES", 1)
    add_numbered(
        doc,
        [
            "Completar diagramas BPMN, UML y DER antes de la entrega final.",
            "Ejecutar y registrar pruebas finales con evidencias.",
            "Cerrar o declarar formalmente los RF pendientes y parciales.",
            "Ampliar la automatizacion E2E para flujos criticos.",
            "Documentar limitaciones del ML actual y plan de evolucion futura.",
        ],
    )
    add_heading(doc, "REFERENCIAS", 1)
    refs = [
        "Carhuamaca, D., & Chuchon, E. (2026). Documento de Requerimientos de Software SIDERAE-Blenkir, version 1.0.",
        "Documentacion oficial de Laravel. https://laravel.com/docs",
        "Documentacion oficial de React. https://react.dev",
        "Documentacion oficial de Docker. https://docs.docker.com",
        "Spatie Laravel Permission. https://spatie.be/docs/laravel-permission",
        "ISO/IEC 25010. Calidad del producto software. Referencia orientativa.",
        "ISO/IEC 29119. Pruebas de software. Referencia orientativa.",
        "Scrum Guide 2020. Referencia metodologica complementaria.",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Bullet")
    add_placeholder(doc, "Revisar formato APA y completar fuentes academicas finales.")
    add_heading(doc, "ANEXOS", 1)
    add_bullets(
        doc,
        [
            "Anexo A: Mockups de interfaz.",
            "Anexo B: Capturas reales del sistema.",
            "Anexo C: Resultados de pruebas automatizadas.",
            "Anexo D: Matriz RF-Sprint-Prueba.",
            "Anexo E: Configuracion Docker y variables de entorno.",
        ],
    )
    add_placeholder(doc, "Insertar anexos finales con evidencias.")


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    cover(doc)
    toc(doc)
    for fn in [
        chapter_1,
        chapter_2,
        chapter_3,
        chapter_4,
        chapter_5,
        chapter_6,
        chapter_7,
        chapter_8,
        chapter_9,
        chapter_10,
        chapter_11,
        chapter_12,
        chapter_13,
        chapter_14,
        closing,
    ]:
        fn(doc)
        if fn is not closing:
            doc.add_page_break()
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
